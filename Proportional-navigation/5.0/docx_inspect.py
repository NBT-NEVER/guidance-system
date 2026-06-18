# _*_coding:UTF-8_*_
"""
开发者: NBT
开发时间: 2026-06-18 00:00:00
功能说明: 检查实验报告 docx 的段落与表格结构，辅助定位需要替换的正文位置
版本号：1.0
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"C:\Users\18030\Desktop\GPT-files\guidance-system\2.0")


def iter_block_items(parent: DocumentObject | _Cell):
    """功能：按文档中的真实顺序遍历段落和表格。参数：父节点。返回：块级对象迭代器。"""
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def main() -> None:
    """功能：打印实验报告中的关键结构信息。参数：无。返回：无。调用位置：python docx_inspect.py。"""
    files = [
        path
        for path in BASE_DIR.glob("*.docx")
        if not path.name.startswith("~$") and "backup" not in path.name.lower()
    ]
    if not files:
        raise FileNotFoundError("未找到目标 docx 文件。")

    doc = Document(str(files[0]))

    print(f"docx={files[0]}")
    print(f"top_level_paragraphs={len(doc.paragraphs)}")
    print(f"top_level_tables={len(doc.tables)}")

    block_index = 0
    for block in iter_block_items(doc):
        block_index += 1
        if isinstance(block, Paragraph):
            text = block.text.strip().replace("\r", "").replace("\n", "")
            if text:
                print(f"[B{block_index:03d}] P style={block.style.name!r} text={text!r}")
        else:
            print(f"[B{block_index:03d}] T rows={len(block.rows)} cols={len(block.columns)}")
            for row_idx, row in enumerate(block.rows, start=1):
                for col_idx, cell in enumerate(row.cells, start=1):
                    for para_idx, para in enumerate(cell.paragraphs, start=1):
                        text = para.text.strip().replace("\r", "").replace("\n", "")
                        if text:
                            print(
                                f"    [T{block_index:03d} R{row_idx}C{col_idx} P{para_idx}] "
                                f"style={para.style.name!r} text={text!r}"
                            )


if __name__ == "__main__":
    main()
