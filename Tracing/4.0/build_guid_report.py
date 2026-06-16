# _*_coding:UTF-8_*_
"""
开发者: OpenAI Codex
文件名: build_guid_report.py
生成时间: 2026-06-11 00:00:00
文件名: build_guid_report.py
功能说明: 按 demo 模板生成追踪法实验报告 / Build the tracing experiment report from the demo template.
"""

from __future__ import annotations

import csv
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


ROOT_DIR = Path(__file__).resolve().parent
SUMMARY_CSV = ROOT_DIR / "outputs" / "tables" / "summary.csv"
POLAR_FIG = ROOT_DIR / "outputs" / "figures" / "attack_zone.png"
FEASIBLE_FIG = ROOT_DIR / "outputs" / "figures" / "attack_zone_cartesian_feasible.png"
INTERCEPTED_FIG = ROOT_DIR / "outputs" / "figures" / "attack_zone_cartesian_intercepted.png"

GUID_DIR = Path(r"C:\Users\18030\Desktop\GPT-files\guid")
TEMPLATE_DOCX = GUID_DIR / "_template_docx" / "demo.docx"
OUTPUT_DOCX = GUID_DIR / "追踪法允许攻击区实验报告.docx"


def set_run_font(
    run: Run,
    chinese_font: str,
    size_pt: float,
    *,
    bold: bool = False,
    ascii_font: str = "Times New Roman",
) -> None:
    """
    功能：统一设置中英文字体、字号和粗细。
    参数：run 为目标文本片段，chinese_font 为中文字体，size_pt 为字号，bold 表示是否加粗，ascii_font 为西文字体。
    返回：无。
    调用位置：正文、封面、表格和标题格式化阶段。
    """

    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def clear_paragraph(paragraph: Paragraph) -> None:
    """
    功能：清空段落文字但保留段落属性。
    参数：paragraph 为待清空段落。
    返回：无。
    调用位置：模板清理和文本替换阶段。
    """

    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    chinese_font: str = "宋体",
    size_pt: float = 12.0,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    """
    功能：重写段落文本并统一设置显示格式。
    参数：paragraph 为目标段落，text 为目标文本，chinese_font 为中文字体，size_pt 为字号，bold 表示是否加粗，alignment 为对齐方式。
    返回：无。
    调用位置：封面与目录内容写入阶段。
    """

    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, chinese_font, size_pt, bold=bold)
    if alignment is not None:
        paragraph.alignment = alignment


def format_document_styles(document: DocxDocument) -> None:
    """
    功能：设置正文与标题的统一样式。
    参数：document 为当前文档对象。
    返回：无。
    调用位置：build_report() 内部。
    """

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading2.font.size = Pt(15)
    heading2.font.bold = True

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Times New Roman"
    heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    heading3.font.size = Pt(14)
    heading3.font.bold = True


def remove_body_content_after_toc(document: DocxDocument, toc_break_paragraph: Paragraph) -> None:
    """
    功能：删除模板中目录之后的全部正文内容，但保留最后的分节属性。
    参数：document 为当前文档对象，toc_break_paragraph 为目录节最后一个分节段落。
    返回：无。
    调用位置：build_report() 内部。
    """

    body = document._element.body
    anchor = toc_break_paragraph._element
    started = False
    for child in list(body):
        if child is anchor:
            started = True
            continue
        if started and child.tag != qn("w:sectPr"):
            body.remove(child)


def format_body_paragraph(paragraph: Paragraph) -> None:
    """
    功能：设置正文普通段落格式。
    参数：paragraph 为目标段落。
    返回：无。
    调用位置：正文段落创建阶段。
    """

    paragraph.style = "Normal"
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_body_paragraph(document: DocxDocument, text: str) -> None:
    """
    功能：追加一个正文普通段落。
    参数：document 为当前文档对象，text 为段落文本。
    返回：无。
    调用位置：正文内容写入阶段。
    """

    paragraph = document.add_paragraph()
    format_body_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12.0)


def add_heading(document: DocxDocument, level: int, text: str) -> None:
    """
    功能：追加指定层级标题。
    参数：document 为当前文档对象，level 为标题层级，text 为标题文本。
    返回：无。
    调用位置：章节结构写入阶段。
    """

    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    paragraph = document.add_paragraph(style=style_name)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", 16.0, bold=True)
    elif level == 2:
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", 15.0, bold=True)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", 14.0, bold=True)


def add_formula(document: DocxDocument, text: str) -> None:
    """
    功能：追加一个居中的公式文本段落。
    参数：document 为当前文档对象，text 为公式文本。
    返回：无。
    调用位置：原理和积分公式说明阶段。
    """

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, "Cambria Math", 12.0, ascii_font="Cambria Math")


def add_caption(document: DocxDocument, text: str, *, above: bool) -> None:
    """
    功能：追加图题或表题。
    参数：document 为当前文档对象，text 为标题文本，above 表示是否位于对象上方。
    返回：无。
    调用位置：表格和图片说明阶段。
    """

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6 if above else 3)
    paragraph.paragraph_format.space_after = Pt(3 if above else 6)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12.0, bold=True)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    """
    功能：统一设置表格单元格文字。
    参数：cell 为目标单元格，text 为文本内容，bold 表示是否加粗。
    返回：无。
    调用位置：表格构建阶段。
    """

    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table: Table, widths_cm: list[float]) -> None:
    """
    功能：设置表格各列宽度。
    参数：table 为目标表格，widths_cm 为列宽列表，单位为 cm。
    返回：无。
    调用位置：表格构建阶段。
    """

    table.autofit = False
    for row in table.rows:
        for cell, width_cm in zip(row.cells, widths_cm):
            cell.width = Cm(width_cm)


def add_parameter_table(document: DocxDocument) -> None:
    """
    功能：插入实验参数表。
    参数：document 为当前文档对象。
    返回：无。
    调用位置：2.2.1 仿真条件。
    """

    add_caption(document, "表1 追踪法允许攻击区判定的仿真参数设置", above=True)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.8, 11.0])
    set_cell_text(table.rows[0].cells[0], "参数", bold=True)
    set_cell_text(table.rows[0].cells[1], "取值", bold=True)

    rows = [
        ("导弹速度", "1.2Ma，程序中按 340m/s 音速换算为 408m/s。"),
        ("目标速度", "0.8Ma，程序中换算为 272m/s。"),
        ("速度比", "1.5。"),
        ("目标航向", "0°，沿 x 轴正向水平匀速直线运动。"),
        ("命中半径", "5m。"),
        ("最大允许过载", "20g。"),
        ("积分步长", "DT = 0.01s。"),
        ("最大仿真时长", "T_MAX = 80s。"),
        ("初始距离", "500m 到 5000m，间隔 250m，共 19 个距离层。"),
        ("视线角采样", "LOS_SAMPLING_MODE = adaptive_arc，弧向点间距 250m。"),
    ]
    for key, value in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], key)
        set_cell_text(row.cells[1], value)


def add_range_summary_table(document: DocxDocument, range_summaries: list[dict[str, str]]) -> None:
    """
    功能：插入代表性距离层统计表。
    参数：document 为当前文档对象，range_summaries 为距离层统计列表。
    返回：无。
    调用位置：2.2.3 实验结果与结果分析。
    """

    add_caption(document, "表2 不同初始距离下的代表性统计结果", above=True)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_widths(table, [2.2, 2.0, 2.0, 2.0, 4.0])
    headers = ["初始距离/m", "总点数", "命中点数", "允许点数", "非180°最大允许角/°"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    for item in range_summaries:
        row = table.add_row()
        set_cell_text(row.cells[0], item["range"])
        set_cell_text(row.cells[1], item["total"])
        set_cell_text(row.cells[2], item["intercepted"])
        set_cell_text(row.cells[3], item["feasible"])
        set_cell_text(row.cells[4], item["max_angle"])


def add_file_role_table(document: DocxDocument) -> None:
    """
    功能：插入代码文件与输出结果说明表。
    参数：document 为当前文档对象。
    返回：无。
    调用位置：程序实现说明部分。
    """

    add_caption(document, "表3 程序文件与输出结果的对应关系", above=True)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [3.0, 3.8, 5.0])
    headers = ["文件", "作用", "主要内容"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    rows = [
        ("config.py", "参数与路径管理", "统一定义速度、过载、步长、采样方式和输出路径。"),
        ("main.py", "主计算入口", "完成网格扫描、纯追踪控制、RK4 积分、结果汇总和三张图输出。"),
        ("summary.csv", "结果统计表", "逐点记录最大需求过载、实际过载、脱靶量、命中和可行性判定。"),
        ("attack_zone.png", "极坐标结果图", "展示初始视线角-初始距离平面中的允许攻击区。"),
        ("attack_zone_cartesian_feasible.png", "直角坐标允许攻击区图", "展示映射到直角坐标后的可行区域形状。"),
        ("attack_zone_cartesian_intercepted.png", "直角坐标命中区域图", "展示限幅积分条件下真实进入命中半径的区域。"),
    ]
    for file_name, role, detail in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], file_name)
        set_cell_text(row.cells[1], role)
        set_cell_text(row.cells[2], detail)


def add_picture_block(document: DocxDocument, image_path: Path, caption: str, width_cm: float) -> None:
    """
    功能：插入图片和对应图题。
    参数：document 为当前文档对象，image_path 为图片路径，caption 为图题，width_cm 为图片宽度。
    返回：无。
    调用位置：实验结果图插入阶段。
    """

    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(3)
    picture_paragraph.paragraph_format.space_after = Pt(3)
    picture_paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(document, caption, above=False)


def load_statistics() -> tuple[dict[str, float | int], list[dict[str, str]]]:
    """
    功能：读取 summary.csv 并计算报告所需统计量。
    参数：无。
    返回：总体统计字典与代表性距离层统计列表。
    调用位置：build_report() 内部。
    """

    with open(SUMMARY_CSV, "r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))

    feasible_rows = [row for row in rows if row["feasible"] == "True"]
    intercepted_rows = [row for row in rows if row["intercepted"] == "True"]
    non_feasible_rows = [row for row in rows if row["feasible"] != "True"]

    range_groups: dict[float, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        range_groups[float(row["initial_range_m"])].append(row)

    selected_ranges = [500.0, 1500.0, 2500.0, 3500.0, 5000.0]
    range_summaries: list[dict[str, str]] = []
    for current_range in selected_ranges:
        group = range_groups[current_range]
        feasible_non180 = sorted(
            float(row["initial_los_deg"])
            for row in group
            if row["feasible"] == "True" and abs(float(row["initial_los_deg"]) - 180.0) > 1e-6
        )
        range_summaries.append(
            {
                "range": f"{current_range:.0f}",
                "total": str(len(group)),
                "intercepted": str(sum(row["intercepted"] == "True" for row in group)),
                "feasible": str(sum(row["feasible"] == "True" for row in group)),
                "max_angle": f"{feasible_non180[-1]:.3f}",
            }
        )

    extra_intercepted_points = [
        (
            float(row["initial_range_m"]),
            float(row["initial_los_deg"]),
        )
        for row in rows
        if row["intercepted"] == "True" and row["feasible"] != "True"
    ]

    statistics = {
        "total_points": len(rows),
        "feasible_count": len(feasible_rows),
        "intercepted_count": len(intercepted_rows),
        "feasible_ratio": len(feasible_rows) / len(rows),
        "intercepted_ratio": len(intercepted_rows) / len(rows),
        "max_required_g": max(float(row["max_required_overload_g"]) for row in rows),
        "max_actual_g": max(float(row["max_actual_overload_g"]) for row in rows),
        "avg_flight_feasible": sum(float(row["flight_time_s"]) for row in feasible_rows) / len(feasible_rows),
        "avg_flight_nonfeasible": sum(float(row["flight_time_s"]) for row in non_feasible_rows) / len(non_feasible_rows),
        "avg_sat_feasible": sum(float(row["saturation_ratio"]) for row in feasible_rows) / len(feasible_rows),
        "avg_sat_nonfeasible": sum(float(row["saturation_ratio"]) for row in non_feasible_rows) / len(non_feasible_rows),
        "avg_miss_feasible": sum(float(row["miss_distance_m"]) for row in feasible_rows) / len(feasible_rows),
        "avg_miss_nonfeasible": sum(float(row["miss_distance_m"]) for row in non_feasible_rows) / len(non_feasible_rows),
        "extra_intercepted_count": len(extra_intercepted_points),
        "extra_intercepted_ranges": "、".join(str(int(item[0])) for item in extra_intercepted_points),
        "extra_intercepted_angles": "、".join(f"{item[1]:.1f}" for item in extra_intercepted_points),
    }
    return statistics, range_summaries


def build_report() -> None:
    """
    功能：根据模板、代码和输出结果生成最终实验报告。
    参数：无。
    返回：无。
    调用位置：main() 内部。
    """

    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(f"未找到模板文件：{TEMPLATE_DOCX}")

    statistics, range_summaries = load_statistics()
    document = Document(str(TEMPLATE_DOCX))
    format_document_styles(document)

    paragraphs = list(document.paragraphs)
    remove_body_content_after_toc(document, paragraphs[46])

    set_paragraph_text(paragraphs[12], "题目：追踪法允许攻击区判定与结果分析", chinese_font="黑体", size_pt=18.0, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(paragraphs[14], "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(paragraphs[31], "2026 年 6 月 11 日", chinese_font="黑体", size_pt=16.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(paragraphs[35], "", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(document, 1, "1 概述")
    add_body_paragraph(
        document,
        "本报告对应 Tracing/4.0 版本的“追踪法允许攻击区判定”程序。实验任务是在二维平面内建立导弹与目标的相对运动模型，在给定速度比和最大法向过载的条件下，扫描不同初始弹目距离和初始视线角，判断目标是否能够被拦截，并在此基础上绘出允许攻击区。代码主体由 config.py、main.py 和 README 说明文档构成，输出结果统一写入 outputs 目录[1-3]。"
    )
    add_body_paragraph(
        document,
        "当前程序采用导弹速度 1.2Ma、目标速度 0.8Ma 的组合，速度比为 1.5；目标始终沿 x 轴方向做匀速直线运动，导弹最大允许过载为 20g，命中半径取 5m，初始距离从 500m 扫描到 5000m。为避免外围圆弧采样过稀，视线角由 adaptive_arc 模式按弧长自适应生成，因此半径越大，采样点数越多。"
    )
    add_body_paragraph(
        document,
        f"本次计算共得到 {statistics['total_points']} 个真实仿真点，其中 {statistics['feasible_count']} 个点满足允许攻击区判据，占总点数的 {statistics['feasible_ratio']:.3%}；{statistics['intercepted_count']} 个点在限幅积分条件下进入命中半径，占总点数的 {statistics['intercepted_ratio']:.3%}。程序同时输出极坐标攻击区图、直角坐标允许攻击区图、直角坐标命中区域图以及 summary.csv 统计表[4-7]。"
    )

    add_heading(document, 1, "2 追踪法允许攻击区判定")
    add_heading(document, 2, "2.1 追踪法导引原理")
    add_heading(document, 3, "2.1.1 相对运动模型")
    add_body_paragraph(
        document,
        "程序采用二维平面坐标系描述弹目相对运动，导弹初始位置固定在原点，目标初始位置由初始弹目距离 R0 和初始视线角 q0 决定。若目标初始坐标记为 (xT, yT)，则程序先将视线角转换为弧度，再由极坐标关系得到目标初始位置。计算过程中每个离散时刻都重新求取弹目相对位移和相对距离，用于命中判定与最小脱靶量统计[1]。"
    )
    add_formula(document, "x_T(0) = R_0 cos q_0，    y_T(0) = R_0 sin q_0")
    add_formula(document, "R = sqrt((x_T - x_M)^2 + (y_T - y_M)^2)")
    add_body_paragraph(
        document,
        "目标在整个仿真过程中保持固定航向 TARGET_HEADING_DEG，位置更新由目标速度和航向共同决定。导弹与目标的相对距离 R 一旦小于命中半径 5m，就认为该组初始条件完成拦截；若未进入命中半径，程序继续记录最近接距离 min_range，将其作为脱靶量。"
    )

    add_heading(document, 3, "2.1.2 控制律与可行性判据")
    add_body_paragraph(
        document,
        "本实验采用纯追踪法，导弹期望航向始终指向目标当前所在位置。程序先由 atan2 计算当前视线角，再用它与导弹当前航向之差构造航向误差，并将该误差限制到 [-π, π]，保证导弹总沿最短方向修正。此后用一步离散近似得到指令航向角速度，再换算为需求法向过载。"
    )
    add_formula(document, "gamma_c = atan2(y_T - y_M，x_T - x_M)")
    add_formula(document, "gamma_dot_cmd = heading_error / DT")
    add_formula(document, "n_req = V_M * gamma_dot_cmd / g_0")
    add_formula(document, "n_act = sat(n_req，-n_max，n_max)")
    add_body_paragraph(
        document,
        "纯追踪法的核心约束在于导弹虽然总是想朝视线方向转弯，但实际过载只能在 ±20g 范围内输出。程序一方面统计未限幅前的最大需求过载 max_required_g，另一方面用限幅后的 actual_g 推进轨迹。最终可行性并不单由命中决定，而是同时要求命中成立且理论需求过载不超过上限。"
    )
    add_formula(document, "feasible = intercepted and (max_required_g <= MAX_OVERLOAD_G)")

    add_heading(document, 2, "2.2 实验结果及分析")
    add_heading(document, 3, "2.2.1 仿真条件")
    add_body_paragraph(
        document,
        "仿真条件由 config.py 统一管理，既包含运动参数，也包含数值积分与采样参数。与攻击区形状最相关的量包括速度比、最大允许过载、积分步长、最大仿真时间以及视线角采样方式。当前版本使用 adaptive_arc 模式，使不同半径圆环上的采样点沿弧向保持近似一致的空间间距。"
    )
    add_parameter_table(document)
    add_body_paragraph(
        document,
        "由表1可以看出，初始距离扫描采用 19 个距离层，但每个距离层上的视线角点数并不相同。500m 距离层仅有 8 个点，5000m 距离层已经扩展到 64 个点。这种设置让外圈边界显示更平滑，同时保持图上的每一个点都对应一次真实仿真，而不是绘图阶段的人工补点。"
    )

    add_heading(document, 3, "2.2.2 程序结构与数值积分")
    add_body_paragraph(
        document,
        "程序实现集中在 main.py 中。外层循环按初始距离遍历 RANGE_VALUES_M，内层循环根据当前半径调用 build_los_values_deg 生成该圆环上的视线角采样点；每一个采样点都独立完成命中判定、过载统计、飞行时间记录以及图表汇总。计算结果写入 summary.csv 和 summary.json，绘图函数再从结果表中直接生成极坐标与直角坐标分布图。"
    )
    add_body_paragraph(
        document,
        "当前版本已经将原先的一阶欧拉更新替换为四阶 Runge-Kutta（RK4）方法。积分状态统一写成 (gamma, xM, yM, xT, yT)，当前离散步内将 actual_g 视作常值，先求状态导数，再按四次采样点做加权平均。这样处理后，导弹转弯段的轨迹更新会比直接用一步直线近似更平滑，最近接距离和飞行时间的离散误差也更小[1][9][10]。"
    )
    add_formula(document, "k_1 = f(t_n，y_n)")
    add_formula(document, "k_2 = f(t_n + DT/2，y_n + DT·k_1/2)")
    add_formula(document, "k_3 = f(t_n + DT/2，y_n + DT·k_2/2)")
    add_formula(document, "k_4 = f(t_n + DT，y_n + DT·k_3)")
    add_formula(document, "y_(n+1) = y_n + DT·(k_1 + 2k_2 + 2k_3 + k_4)/6")

    add_heading(document, 3, "2.2.3 实验结果与结果分析")
    add_body_paragraph(
        document,
        "极坐标允许攻击区图直接展示了初始视线角和初始距离对拦截结果的影响。绿色点表示满足允许攻击区判据的初始点，红色点表示命中失败或者需求过载超限的初始点。图中最明显的特征是：当初始视线角较小或者接近 180° 时，导弹更容易进入命中半径；当视线角逐步增大到侧向区间后，导弹需要在末段快速掉头，需求过载会迅速抬升，禁区因此形成[4][8]。"
    )
    add_picture_block(document, POLAR_FIG, "图1 追踪法允许攻击区的极坐标分布", 13.5)
    add_range_summary_table(document, range_summaries)
    add_body_paragraph(
        document,
        "由表2可以看到，随着初始距离增大，单个圆环上的采样点数明显增加，允许点数也整体增多。若将 180° 这一对称特例单独拿开，只看普通角度段的上边界，500m 距离层的最大允许角只有 51.429°，1500m 时增大到 94.737°，5000m 时已经达到 131.429°。这说明在较大初始距离下，导弹有更长的调整路径，能够在不违反过载约束的情况下逐步转入有利攻角。"
    )
    add_body_paragraph(
        document,
        f"从总体统计看，允许点平均飞行时间约为 {statistics['avg_flight_feasible']:.3f}s，平均过载饱和比例为 {statistics['avg_sat_feasible']:.3f}，平均脱靶量约为 {statistics['avg_miss_feasible']:.3f}m；非允许点平均飞行时间只有 {statistics['avg_flight_nonfeasible']:.3f}s，但平均过载饱和比例升到 {statistics['avg_sat_nonfeasible']:.3f}，平均脱靶量达到 {statistics['avg_miss_nonfeasible']:.3f}m。程序统计到的最大实际过载始终被限制在 {statistics['max_actual_g']:.1f}g，而最大理论需求过载达到 {statistics['max_required_g']:.3f}g，这个数值非常大，反映的不是导弹真实输出，而是某些禁区点在当前控制律下对瞬时转向能力提出了远超 20g 的要求。"
    )

    add_heading(document, 2, "2.3 小结")
    add_body_paragraph(
        document,
        "本章工作给出了纯追踪法允许攻击区的判定流程，并把速度条件、过载限制、积分方法和采样方式统一到了同一套程序里。实验结果说明，攻击区的外形并不是单由速度比决定，视线角几何关系和末段转弯能力同样起决定作用。"
    )

    add_heading(document, 1, "3 直角坐标允许攻击区分布")
    add_body_paragraph(
        document,
        "将极坐标初始点映射到直角坐标后，允许攻击区的几何形状更直观。程序在绘图时按 x 轴做镜像扩展，因此上半平面和下半平面的点完全对称。蓝色空心圆表示允许攻击区，黑色方块表示导弹初始位置。"
    )
    add_picture_block(document, FEASIBLE_FIG, "图2 直角坐标下的允许攻击区分布", 12.5)
    add_body_paragraph(
        document,
        "图2中可以看到，允许攻击区主要分布在导弹前方及其附近的宽扇形区域，随着初始距离增大，蓝圈向外扩展得更明显。位于目标侧后方的大角度区域更容易形成禁区，因为纯追踪法总是沿目标当前位置修正航向，导弹在末段需要做更剧烈的横向转弯。图中 x 轴负向保留了一条靠近 180° 的可行支路，这一支路对应目标正向迎面飞来的情况，导弹虽然初始视线角大，但几何关系反而更有利。"
    )

    add_heading(document, 1, "4 命中区域与允许攻击区对比")
    add_body_paragraph(
        document,
        "命中区域图与允许攻击区图的差别集中反映了 intercepted 和 feasible 两个判据的不同。intercepted 只要求弹目距离进入命中半径，feasible 还要求理论需求过载不超过 20g，因此命中区域通常略大于允许攻击区。"
    )
    add_picture_block(document, INTERCEPTED_FIG, "图3 直角坐标下的命中区域分布", 12.5)
    add_body_paragraph(
        document,
        f"本次仿真中共有 {statistics['extra_intercepted_count']} 个点出现“已命中但不计入允许攻击区”的情况，它们主要落在 {statistics['extra_intercepted_ranges']}m 这些距离层上，对应的视线角大致集中在 {statistics['extra_intercepted_angles']}° 一带。这说明在限幅后的真实积分轨迹中，导弹仍有机会擦入 5m 命中半径，但从设计约束角度看，这些点已经需要超过 20g 的理论转向能力，因此不宜算作允许攻击区。命中区域和允许攻击区之间的这条窄带，正好体现了“能碰到”和“满足设计裕度”之间的区别。"
    )

    add_heading(document, 1, "5 程序实现说明")
    add_body_paragraph(
        document,
        "从工程实现角度看，config.py 负责统一管理所有路径与参数，main.py 则承担计算入口、采样生成、积分推进、结果写表和图形输出等功能。这样的分工比较清楚：当实验条件变化时，优先修改 config.py；当算法逻辑变化时，则集中在 main.py 内部调整。"
    )
    add_file_role_table(document)
    add_body_paragraph(
        document,
        "当前版本还有两个实现细节值得保留。其一，build_los_values_deg 同时保留了 uniform_angle 与 adaptive_arc 两种模式，便于后续对比“真实采样点”与“视觉补点”的差异；其二，plot_cartesian_attack_zone 同时输出允许攻击区图和命中区域图，避免在解释实验结果时把命中判定与可行判定混为一谈。"
    )

    add_heading(document, 1, "6 结果讨论")
    add_body_paragraph(
        document,
        "这份程序已经能较完整地反映纯追踪法的几何特点，但它仍然是一个简化模型。目标航向固定不变，没有加入机动；控制律采用 gamma_dot_cmd = heading_error / DT 的离散近似，因此需求过载对步长比较敏感；轨迹终止条件中还使用了 closing_speed > 0 且当前距离明显大于最小距离的启发式判断。RK4 已经改善了积分精度，但并没有从根本上改变控制律的离散近似特征。"
    )
    add_body_paragraph(
        document,
        "如果后续继续完善，可以从三个方向着手：一是把命中时刻和最小距离从离散点判定改成步间插值；二是把纯追踪法与比例导引等导引律放在同一组参数下比较；三是在目标模型中加入转弯或速度变化，考察攻击区是否仍保持当前这种对称形状。这样做以后，实验报告里的攻击区边界会更稳，结果解释也会更接近实际导引问题。"
    )

    add_heading(document, 1, "总结")
    add_body_paragraph(
        document,
        "本次实验围绕追踪法允许攻击区的确定展开，完成了二维相对运动建模、纯追踪控制律实现、最大过载限制、RK4 数值积分、极坐标与直角坐标结果展示等工作。程序输出说明，在速度比固定为 1.5、最大允许过载为 20g 的条件下，攻击区边界主要由视线角几何和末段转弯能力共同决定。"
    )
    add_body_paragraph(
        document,
        "从结果图和统计表可以看出，纯追踪法在小视线角和迎面几何条件下更容易命中，而侧向大视线角区域容易形成禁区。允许攻击区与命中区域并不完全重合，这一差别提醒我们：导引算法是否“能打中”与是否“满足设计约束”并不是同一个问题。"
    )

    add_heading(document, 1, "参考文献")
    references = [
        "[1] NBT. 追踪法允许攻击区分析主程序 main.py[Z]. 2026.",
        "[2] NBT. 追踪法允许攻击区分析配置文件 config.py[Z]. 2026.",
        "[3] NBT. 实验 4.0：追踪法允许攻击区判定 README[Z]. 2026.",
        "[4] NBT. attack_zone.png 允许攻击区极坐标散点图[Z]. 2026.",
        "[5] NBT. attack_zone_cartesian_feasible.png 直角坐标允许攻击区分布图[Z]. 2026.",
        "[6] NBT. attack_zone_cartesian_intercepted.png 直角坐标命中区域分布图[Z]. 2026.",
        "[7] NBT. summary.csv 实验结果统计表[Z]. 2026.",
        "[8] Zarchan P. Tactical and Strategic Missile Guidance[M]. American Institute of Aeronautics and Astronautics.",
        "[9] Burden R L, Faires J D. Numerical Analysis[M]. Brooks/Cole.",
        "[10] Chapra S C, Canale R P. Numerical Methods for Engineers[M]. McGraw-Hill.",
    ]
    for item in references:
        paragraph = document.add_paragraph()
        paragraph.style = "Normal"
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(item)
        set_run_font(run, "宋体", 12.0)

    document.save(str(OUTPUT_DOCX))


def main() -> None:
    """
    功能：执行实验报告生成流程。
    参数：无。
    返回：无。
    调用位置：python build_guid_report.py。
    """

    build_report()


if __name__ == "__main__":
    main()
