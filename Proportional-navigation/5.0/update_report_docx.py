# _*_coding:UTF-8_*_
"""
开发者: NBT
开发时间: 2026-06-18 00:00:00
功能说明: 定点修改比例导引法允许攻击区实验报告中与四种最大可用过载对比相关的正文分析
版本号：1.0
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


DOCX_PATH = Path(r"C:\Users\18030\Desktop\GPT-files\guidance-system\2.0\比例导引法允许攻击区实验报告.docx")
BACKUP_PATH = Path(
    r"C:\Users\18030\Desktop\GPT-files\guidance-system\2.0\比例导引法允许攻击区实验报告.backup_before_overload_analysis.docx"
)


PARAGRAPH_REPLACEMENTS = [
    (
        "为了更直观地观察允许攻击区在平面中的位置关系",
        "为了比较最大可用过载对允许攻击区的影响，本节分别给出最大可用过载为 5 g、10 g、15 g 和 20 g 时的直角坐标允许攻击区分布。图中黑色方块表示目标初始位置，红色散点表示满足脱靶量小于 0.1 m 的初始发射点。四幅图采用相同坐标范围，因此四种过载条件下的攻击区边界能够直接比较。",
    ),
    (
        "图 3-1 中，允许攻击区主要分布在目标左后方",
        "图 3-1 至图 3-4 反映出最大可用过载增大后允许攻击区持续外扩，但外扩方向并不均匀。5 g 时，可行发射点主要集中在目标左后方，目标前方仅保留一条较窄的可行带，说明低过载条件下导弹对大偏角阵位的修正能力有限。过载上限提高到 10 g 和 15 g 后，左后方区域先变得更饱满，靠近目标附近的弧形边界向原点收缩，右前方的细长可行带逐步扩展为扇形区域。20 g 时，右侧与侧前方允许攻击区继续明显展开，上下边界张开幅度最大，目标前方大部分远距阵位已经转为可行；但目标附近仍保留小块不可行区，这说明允许攻击区不仅受过载上限控制，还与初始距离过短、视线角变化过快和可调整时间不足有关。",
    ),
    ("图 2 比例导引法允许攻击区参数平面分布图", "图 4-1 比例导引法允许攻击区参数平面分布图"),
    (
        "图 2 中最清楚的现象是“半径越大，可行角范围越宽”",
        "图 4-1 仍然表明发射半径越大，可行角范围越宽：125 m 半径层仅保留 0° 和 180° 两个点，500 m 半径层扩展到 195°，1500 m 以后已经覆盖到 355°。把该图与第 3 节四种最大可用过载的直角坐标结果对照，参数平面中的外扩主要对应目标前方和侧前方阵位在提高过载上限后的逐步放宽。0°附近若干中等距离层仍然保留高需用过载样本，说明拦截难点并不只由距离决定。",
    ),
    (
        "本次结果最突出的现象不是“导弹速度高于目标速度”",
        "结合第 3 节的四幅直角坐标图，本次结果中最突出的现象不是“导弹速度高于目标速度”，而是初始几何关系与最大可用过载共同决定了过载需求的增长方式。最大过载由 5 g 提高到 20 g 时，允许攻击区首先在目标前方和侧前方明显展开，左后方区域的变化相对平缓。理论需用过载峰值超过两万 g 的样本仍主要集中在半径约 1125 m 到 1250 m、方位角靠近 5° 和 355° 的区域。这一带并非最远距离点，但它同时具备较高的视线角变化率和较短的可调整时间，因此更容易把比例导引指令推到极端。",
    ),
    (
        "更重要的是，这组结果把“导弹能否命中”从单条弹道问题转成了阵位条件问题。",
        "更重要的是，这组结果把“导弹能否命中”从单条弹道问题转成了阵位条件问题。对当前程序而言，5 g 工况下的过载上限仍是主导约束；当最大可用过载提高到 10 g、15 g 和 20 g 后，目标前方和侧前方的可行阵位连续增加，但目标附近仍保留不可行区。说明视线角变化率、闭合速度和可调整时间仍是决定允许攻击区边界的主要变量。后续若继续修改代码，最值得优先检验的仍是过载上限、积分方式和最小距离估计是否发生了实质变化。",
    ),
]


CAPTION_REPLACEMENTS = [
    ("图 3-1 最大可用过载5g允许攻击区", "图 3-1 最大可用过载为 5 g 时的允许攻击区"),
    ("图 3-2 最大可用过载10g允许攻击区", "图 3-2 最大可用过载为 10 g 时的允许攻击区"),
    ("图 3-3 最大可用过载15g允许攻击区", "图 3-3 最大可用过载为 15 g 时的允许攻击区"),
    ("图 3-4 最大可用过载20g允许攻击区", "图 3-4 最大可用过载为 20 g 时的允许攻击区"),
]


def replace_text_in_paragraphs(doc: Document) -> int:
    """功能：替换正文段落内容。参数：doc。返回：替换数量。"""
    count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for needle, replacement in PARAGRAPH_REPLACEMENTS:
            if needle in text:
                paragraph.text = replacement
                count += 1
                break
    return count


def replace_text_in_tables(doc: Document) -> int:
    """功能：替换表格单元格中的图注内容。参数：doc。返回：替换数量。"""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for needle, replacement in CAPTION_REPLACEMENTS:
                        if needle in text:
                            paragraph.text = replacement
                            count += 1
                            break
    return count


def main() -> None:
    """功能：备份并更新目标报告。参数：无。返回：无。调用位置：python update_report_docx.py。"""
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"未找到目标文档：{DOCX_PATH}")

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    paragraph_count = replace_text_in_paragraphs(doc)
    caption_count = replace_text_in_tables(doc)
    doc.save(str(DOCX_PATH))

    print(f"paragraph_replacements={paragraph_count}")
    print(f"caption_replacements={caption_count}")
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
